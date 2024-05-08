from odoo import models, fields
import docx
import base64
from io import BytesIO
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.dml import MSO_COLOR_TYPE
from docx.shared import RGBColor, Pt
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls


class TemporaryDocument(models.Model):
    _name = 'temporary.document'
    _description = 'Temporary Document'
    
    
    # File name
    name = fields.Char(string='Name')
    description = fields.Text('Description')
    
    # Purpose
    purpose = fields.Text('Purpose')
    
    # Require
    require = fields.Text('Require')
    
    # Time
    time = fields.Text('Time')
    
    # Place
    place = fields.Text('Place')
    
    # Object
    object_participant = fields.Text('Object Participant')
    
    # Number of participant
    num_of_participant = fields.Text('Number Of Participant')
    
    # Content
    content = fields.Text('Content')
    
    word_file = fields.Binary('Word File')
    word_filename = fields.Char('Filename')
    
    def generate_word_file(self):
        # Create a new Word document
        doc = docx.Document()
        
        # Add header
        table = doc.add_table(rows=2, cols=2)
        row_first = table.rows[0].cells
        row_first[0].text = 'ĐOÀN TRƯỜNG ĐẠI HỌC BÁCH KHOA'
        row_first[1].text = 'ĐOÀN TNCS HỒ CHÍ MINH'
        row_second = table.rows[1].cells
        row_second[0].text = 'BCH ĐOÀN KHOA KH&KT MÁY TÍNH'
        
        for cell in table._cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.color.rgb = RGBColor(0, 0, 0)  # Set font color to black
                    run.font.bold = True  # Remove bold formatting
                    
        # Adjust the width of the first column
        first_column = table.columns[0]
        first_column.width = Pt(270)  # Set the desired width (adjust as needed)
        
        # Set white border for each cell
        for row in table.rows:
            for cell in row.cells:
                tcPr = cell._tc.get_or_add_tcPr()
                tcBorders = parse_xml(r'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                                    r'<w:top w:val="single" w:color="FFFFFF"/>'
                                    r'<w:bottom w:val="single" w:color="FFFFFF"/>'
                                    r'<w:left w:val="single" w:color="FFFFFF"/>'
                                    r'<w:right w:val="single" w:color="FFFFFF"/>'
                                    r'</w:tcBorders>')
                tcPr.append(tcBorders)
        
        
        # Add title
        title = doc.add_heading('KẾ HOẠCH', level=1)
        title_format = title.paragraph_format
        title_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add name
        name = doc.add_heading(str(self.name), level=3)
        name_format = name.paragraph_format
        name_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add content to the document
        doc.add_heading('I. MỤC ĐÍCH - YÊU CẦU', level=2)
        doc.add_heading('1. Mục đích', level=3)
        doc.add_paragraph(str(self.purpose))
        doc.add_heading('2. Yêu cầu', level=3)
        doc.add_paragraph(str(self.require))
        doc.add_heading('II. THỜI GIAN - ĐỊA ĐIỂM TỔ CHỨC', level=2)
        doc.add_heading('1. Thời gian', level=3)
        doc.add_paragraph(str(self.time))
        doc.add_heading('2. Địa điểm', level=3)
        doc.add_paragraph(str(self.place))
        doc.add_heading('III. ĐỐI TƯỢNG THAM GIA', level=2)
        doc.add_paragraph(str(self.object_participant))
        doc.add_heading('IV. NỘI DUNG CHƯƠNG TRÌNH ', level=2)
        doc.add_paragraph(str(self.content))
        
        # Save the document to a BytesIO object
        byte_io = BytesIO()
        doc.save(byte_io)

        # Convert the BytesIO object to a base64 encoded string
        word_file = base64.b64encode(byte_io.getvalue()).decode('utf-8')
        
        # Update the record with the Word file and filename
        self.write({
            'word_file': word_file,
            'word_filename': "document.docx",
        })
        
    def get_document_url(self):
        self.ensure_one()
        self.generate_word_file()
        base_url = self.env['ir.config_parameter'].sudo().get_param('web.base.url')
        url = f"{base_url}/download/document/{self.name}"
        return {
            'type': 'ir.actions.act_url',
            'url': str(url),
            'target': 'new',
        }
